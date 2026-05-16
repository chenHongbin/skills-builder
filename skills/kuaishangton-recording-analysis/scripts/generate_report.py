#!/usr/bin/env python3
"""
快商通录音分析 - Word 报告生成器
将 8 大版块分析结果生成为格式化的 Word 文档。
纯 python-docx 实现。

用法：python3 generate_report.py <分析结果.json>
"""

import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx。运行：pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        set_cell_shading(cell, '2B579A')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()  # 表后空行
    return table


def generate_report(data):
    """根据分析结果 JSON 生成 Word 报告"""
    doc = Document()

    # ============================================================
    # 页面设置
    # ============================================================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ============================================================
    # 页眉：快商通 LOGO + 报告标题
    # ============================================================
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run('快商通 · 电销团队录音分析报告')
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    header_run.font.name = 'Arial'

    # ============================================================
    # 封面标题
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('电销团队录音分析报告')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(43, 87, 154)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sales_name = data.get('sales_name', '未知')
    date_str = data.get('date', datetime.now().strftime('%y.%m.%d'))
    customer_type = data.get('customer_type', '未知')
    subtitle_run = subtitle.add_run(f'{sales_name} · {customer_type}客户 · {date_str}')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 分隔线
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = separator.add_run('━' * 40)
    sep_run.font.color.rgb = RGBColor(43, 87, 154)
    sep_run.font.size = Pt(8)

    doc.add_paragraph()

    # ============================================================
    # 版块一：客户画像
    # ============================================================
    doc.add_heading('一、客户画像（基础信息与现状）', level=1)
    profile = data.get('sections', {}).get('customer_profile', {})

    if profile.get('basic_info'):
        p = doc.add_paragraph()
        run = p.add_run(f"基本信息：{profile['basic_info']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    doc.add_heading('核心现状', level=3)
    for item in profile.get('current_status', []):
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = 'Arial'

    if profile.get('is_key_account'):
        p = doc.add_paragraph()
        run = p.add_run(f"💡 判断：{profile['is_key_account']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.bold = True

    doc.add_paragraph()

    # ============================================================
    # 版块二：通话核心：价值传递，诊断与算账
    # ============================================================
    doc.add_heading('二、通话核心：价值传递，诊断与算账', level=1)
    value_diag = data.get('sections', {}).get('value_diagnosis', {})

    if value_diag:
        headers = ['维度', '分析内容']
        rows = [
            ['价值传递', value_diag.get('value_proposition', '未提及')],
            ['数据算账', value_diag.get('roi_calculation', '❌ 未算账')],
            ['医疗痛点', value_diag.get('pain_points', '未触及')],
        ]
        add_styled_table(doc, headers, rows)

    # ============================================================
    # 版块三：客户意向与信号
    # ============================================================
    doc.add_heading('三、客户意向与信号', level=1)
    intent = data.get('sections', {}).get('customer_intent', {})

    if intent.get('rating'):
        rating_map = {'高': '🟢 高', '中': '🟡 中', '低': '🔴 低'}
        rating_text = rating_map.get(intent['rating'], intent['rating'])
        p = doc.add_paragraph()
        run = p.add_run(f"综合评级：{rating_text}")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Arial'

    doc.add_heading('关键信号', level=3)
    for signal in intent.get('positive_signals', []):
        p = doc.add_paragraph(f'✅ {signal}', style='List Bullet')
    for signal in intent.get('negative_signals', []):
        p = doc.add_paragraph(f'❌ {signal}', style='List Bullet')

    doc.add_paragraph()

    # ============================================================
    # 版块四：话术红黑榜
    # ============================================================
    doc.add_heading('四、话术红黑榜', level=1)
    script_review = data.get('sections', {}).get('script_review', {})

    if script_review.get('good_points'):
        doc.add_heading('✅ 做得好 👍', level=3)
        for point in script_review['good_points']:
            p = doc.add_paragraph(point, style='List Bullet')

    if script_review.get('improvements'):
        doc.add_heading('🔧 待改进', level=3)
        for imp in script_review['improvements']:
            p = doc.add_paragraph()
            run = p.add_run(f"问题：{imp.get('problem', '')}")
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.bold = True
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"建议话术：{imp.get('suggestion', '')}")
            run2.font.size = Pt(11)
            run2.font.name = 'Arial'
            run2.italic = True

    doc.add_paragraph()

    # ============================================================
    # 版块五：竞争情报（竞品攻防）
    # ============================================================
    doc.add_heading('五、竞争情报（竞品攻防）', level=1)
    competitive = data.get('sections', {}).get('competitive_intel', {})

    if competitive.get('current_tool'):
        p = doc.add_paragraph()
        run = p.add_run(f"现用工具：{competitive['current_tool']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    if competitive.get('dissatisfaction'):
        p = doc.add_paragraph()
        run = p.add_run(f"客户不满：{competitive['dissatisfaction']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True

    doc.add_heading('快商通打击点', level=3)
    for point in competitive.get('attack_points', []):
        p = doc.add_paragraph(f'🎯 {point}', style='List Bullet')

    doc.add_paragraph()

    # ============================================================
    # 版块六：24h/3d 关键动作
    # ============================================================
    doc.add_heading('六、24h / 3d 关键动作', level=1)
    actions = data.get('sections', {}).get('action_items', {})

    if actions.get('within_24h'):
        doc.add_heading('⏰ 24小时内', level=3)
        for action in actions['within_24h']:
            p = doc.add_paragraph(action, style='List Bullet')

    if actions.get('within_3d'):
        doc.add_heading('📅 3天内', level=3)
        for action in actions['within_3d']:
            p = doc.add_paragraph(action, style='List Bullet')

    doc.add_paragraph()

    # ============================================================
    # 版块七：经理点评与总结
    # ============================================================
    doc.add_heading('七、经理点评与总结', level=1)
    summary = data.get('sections', {}).get('manager_summary', {})

    if summary.get('closing_probability'):
        p = doc.add_paragraph()
        run = p.add_run(f"成交可能性：{summary['closing_probability']}%")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Arial'

    if summary.get('one_liner'):
        p = doc.add_paragraph()
        run = p.add_run(f"一句话定性：{summary['one_liner']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True

    if summary.get('coaching_priority'):
        p = doc.add_paragraph()
        run = p.add_run(f"辅导重点：{summary['coaching_priority']}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    doc.add_paragraph()

    # ============================================================
    # 版块八：综合对比表（仅多通录音时显示）
    # ============================================================
    comparison = data.get('sections', {}).get('comparison')
    if comparison and comparison.get('recordings'):
        doc.add_heading('八、综合对比表', level=1)
        recordings = comparison['recordings']
        dimensions = comparison.get('dimensions', [
            '开场方式', '价值传递', '竞品处理', '意向挖掘', '下一步动作', '综合得分'
        ])

        headers = ['维度'] + [r['label'] for r in recordings]
        rows = []
        for dim in dimensions:
            row = [dim]
            for rec in recordings:
                row.append(rec.get('scores', {}).get(dim, '-'))
            rows.append(row)

        add_styled_table(doc, headers, rows)

    # ============================================================
    # 页脚
    # ============================================================
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('快商通 · AI 驱动销售增长 | 录音数据不外传，仅用于团队辅导')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.font.name = 'Arial'

    # ============================================================
    # 保存文件
    # ============================================================
    output_path = data.get('output_path', '')
    if not output_path:
        output_path = os.path.expanduser(f'~/Desktop/录音分析/{sales_name}/')

    os.makedirs(output_path, exist_ok=True)

    filename = f'{date_str}_{sales_name}_{customer_type}_录音分析.docx'
    full_path = os.path.join(output_path, filename)

    doc.save(full_path)

    return full_path


def main():
    if len(sys.argv) < 2:
        print('用法: python3 generate_report.py <分析结果.json>', file=sys.stderr)
        print('将分析结果 JSON 生成为格式化的 Word 报告。', file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f'错误: 文件不存在 "{input_path}"', file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f'错误: JSON 格式无效 - {e}', file=sys.stderr)
        sys.exit(1)

    output_path = generate_report(data)
    print(json.dumps({
        'success': True,
        'output_path': output_path,
        'sales_name': data.get('sales_name', '未知'),
        'customer_type': data.get('customer_type', '未知'),
    }, ensure_ascii=False))


if __name__ == '__main__':
    main()
