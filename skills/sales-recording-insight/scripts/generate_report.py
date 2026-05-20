#!/usr/bin/env python3
"""
电销录音管理教练 - Word 报告生成器
将分析结果生成为格式化的 Word 文档。
纯 python-docx 实现。

用法：python3 generate_report.py <分析结果.json>
"""

import json
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx。运行：pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def set_cell_shading(cell, color):
    """设置表格单元格底色"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading_elm.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Arial'
        set_cell_shading(cell, '2B579A')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()  # 表后空行
    return table


def add_bullet_paragraph(doc, text, prefix=""):
    """添加带前缀的项目符号段落"""
    p = doc.add_paragraph(f"{prefix} {text}", style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p


def generate_report(data):
    """根据分析结果 JSON 生成 Word 报告"""
    doc = Document()

    # ============================================================
    # 页面设置
    # ============================================================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ============================================================
    # 页眉：报告标题
    # ============================================================
    header = section.header
    header.is_linked_to_previous = False
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run('电销录音管理教练 · 分析报告')
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(128, 128, 128)
    header_run.font.name = 'Arial'

    # ============================================================
    # 封面标题
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('电销录音管理教练报告')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(43, 87, 154)

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sales_name = data.get('sales_name', '未知')
    date_str = data.get('date', datetime.now().strftime('%y.%m.%d'))
    customer_type = data.get('customer_type', '未知')
    subtitle_run = subtitle.add_run(f'{sales_name} · {customer_type}客户 · {date_str}')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 空行

    # 分隔线
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = separator.add_run('━' * 40)
    sep_run.font.color.rgb = RGBColor(43, 87, 154)
    sep_run.font.size = Pt(8)

    doc.add_paragraph()

    # ============================================================
    # 维度一：通话摘要
    # ============================================================
    doc.add_heading('一、通话摘要', level=1)
    summary = data.get('sections', {}).get('call_summary', {})

    if summary:
        headers = ['维度', '内容']
        rows = [
            ['客户', summary.get('customer', '未知')],
            ['需求', summary.get('need', '未明确提及')],
            ['异议', summary.get('objection', '未提及')],
            ['销售表达', summary.get('sales_pitched', '未明确提及')],
            ['结果', summary.get('outcome', '未明确')],
            ['时长', summary.get('duration', '未知')],
        ]
        add_styled_table(doc, headers, rows)

    doc.add_paragraph()

    # ============================================================
    # 维度二：关键标签
    # ============================================================
    doc.add_heading('二、关键标签', level=1)
    tags = data.get('sections', {}).get('tags', {})

    if tags.get('core'):
        doc.add_heading('核心标签', level=3)
        for tag in tags['core']:
            add_bullet_paragraph(doc, tag, '🔹')

    if tags.get('problems'):
        doc.add_heading('问题标签', level=3)
        for tag in tags['problems']:
            add_bullet_paragraph(doc, tag, '🔴')

    if tags.get('strengths'):
        doc.add_heading('能力标签', level=3)
        for tag in tags['strengths']:
            add_bullet_paragraph(doc, tag, '✅')

    doc.add_paragraph()

    # ============================================================
    # 维度三：话术评分
    # ============================================================
    doc.add_heading('三、话术评分', level=1)
    scores = data.get('sections', {}).get('scores', {})
    overall = data.get('sections', {}).get('overall_score', {})

    if scores:
        headers = ['评分维度', '得分', '评价理由']
        dimension_names = {
            'opening': '开场有效性',
            'need_discovery': '需求挖掘充分度',
            'objection_handling': '异议处理能力',
            'value_clarity': '价值表达清晰度',
            'next_step': '推进下一步明确度',
        }
        rows = []
        for key, name in dimension_names.items():
            item = scores.get(key, {})
            score = item.get('score', '-')
            reason = item.get('reason', '未评价')
            rows.append([name, f'{score}/10', reason])

        add_styled_table(doc, headers, rows)

    # 综合得分
    if overall:
        p = doc.add_paragraph()
        run = p.add_run(f"综合得分：{overall.get('score', '-')}/10")
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = 'Arial'

        if overall.get('calculation'):
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"计算方式：{overall['calculation']}")
            run2.font.size = Pt(10)
            run2.font.name = 'Arial'
            run2.italic = True

        if overall.get('level'):
            p3 = doc.add_paragraph()
            run3 = p3.add_run(f"等级：{overall['level']}")
            run3.font.size = Pt(11)
            run3.font.name = 'Arial'

    doc.add_paragraph()

    # ============================================================
    # 维度四：话术红黑榜
    # ============================================================
    doc.add_heading('四、话术红黑榜', level=1)
    script_review = data.get('sections', {}).get('script_review', {})

    if script_review.get('good_points'):
        doc.add_heading('✅ 亮点（做得好）', level=3)
        for point in script_review['good_points']:
            moment = point.get('moment', '')
            what = point.get('what', '')
            text = f"{moment}：{what}" if moment else what
            add_bullet_paragraph(doc, text, '👍')

    if script_review.get('improvements'):
        doc.add_heading('🔧 待改进', level=3)
        for imp in script_review['improvements']:
            moment = imp.get('moment', '')
            problem = imp.get('problem', '')
            suggestion = imp.get('suggestion', '')

            if moment:
                p = doc.add_paragraph()
                run = p.add_run(f"【{moment}】")
                run.font.size = Pt(11)
                run.font.name = 'Arial'
                run.bold = True

            p = doc.add_paragraph()
            run = p.add_run(f"问题：{problem}")
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            run.bold = True

            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"建议话术：{suggestion}")
            run2.font.size = Pt(11)
            run2.font.name = 'Arial'
            run2.italic = True

            doc.add_paragraph()  # 每条改进之间空一行

    doc.add_paragraph()

    # ============================================================
    # 维度五：改进建议
    # ============================================================
    doc.add_heading('五、改进建议', level=1)
    advice = data.get('sections', {}).get('actionable_advice', {})

    if advice.get('immediate'):
        doc.add_heading('⚡ 立即改进（下次就能用）', level=3)
        imm = advice['immediate']
        if imm.get('action'):
            p = doc.add_paragraph()
            run = p.add_run(f"动作：{imm['action']}")
            run.font.size = Pt(11)
            run.font.name = 'Arial'
        if imm.get('script'):
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"示范话术：{imm['script']}")
            run2.font.size = Pt(11)
            run2.font.name = 'Arial'
            run2.italic = True

    if advice.get('short_term'):
        doc.add_heading('📈 短期提升（本周重点练）', level=3)
        st = advice['short_term']
        if st.get('focus'):
            add_bullet_paragraph(doc, f"重点：{st['focus']}")
        if st.get('practice'):
            add_bullet_paragraph(doc, f"练习方式：{st['practice']}")
        if st.get('success_criteria'):
            add_bullet_paragraph(doc, f"验收标准：{st['success_criteria']}")

    if advice.get('long_term'):
        doc.add_heading('🎯 长期成长（持续优化）', level=3)
        lt = advice['long_term']
        if lt.get('direction'):
            add_bullet_paragraph(doc, f"方向：{lt['direction']}")
        if lt.get('method'):
            add_bullet_paragraph(doc, f"方法：{lt['method']}")

    doc.add_paragraph()

    # ============================================================
    # 维度六：管理洞察
    # ============================================================
    doc.add_heading('六、管理洞察', level=1)
    insight = data.get('sections', {}).get('management_insight', {})

    if insight.get('core_problem'):
        p = doc.add_paragraph()
        run = p.add_run(f"核心问题：{insight['core_problem']}")
        run.font.size = Pt(12)
        run.font.name = 'Arial'
        run.bold = True

    if insight.get('coaching_priority'):
        doc.add_heading('辅导优先级', level=3)
        for i, priority in enumerate(insight['coaching_priority'], 1):
            add_bullet_paragraph(doc, f"{i}. {priority}")

    if insight.get('customer_worth'):
        worth = insight['customer_worth']
        doc.add_heading('客户跟进建议', level=3)
        judgment = worth.get('judgment', '')
        reason = worth.get('reason', '')
        p = doc.add_paragraph()
        run = p.add_run(f"{judgment}")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.bold = True
        if reason:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f"理由：{reason}")
            run2.font.size = Pt(11)
            run2.font.name = 'Arial'

    if insight.get('one_on_one_script'):
        doc.add_heading('一对一辅导话术（给主管参考）', level=3)
        p = doc.add_paragraph()
        run = p.add_run(insight['one_on_one_script'])
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.italic = True

    if insight.get('team_lesson'):
        lesson = insight['team_lesson']
        doc.add_heading('团队经验', level=3)
        if lesson.get('positive'):
            add_bullet_paragraph(doc, lesson['positive'], '👍')
        if lesson.get('warning'):
            add_bullet_paragraph(doc, lesson['warning'], '⚠️')

    doc.add_paragraph()

    # ============================================================
    # 多通录音：团队诊断（如有）
    # ============================================================
    team_diag = data.get('sections', {}).get('team_diagnosis')
    if team_diag:
        doc.add_heading('七、团队诊断', level=1)

        if team_diag.get('strengths'):
            doc.add_heading('团队优势', level=3)
            for item in team_diag['strengths']:
                add_bullet_paragraph(doc, item, '✅')

        if team_diag.get('weaknesses'):
            doc.add_heading('团队短板', level=3)
            for item in team_diag['weaknesses']:
                add_bullet_paragraph(doc, item, '🔴')

        if team_diag.get('training_theme'):
            p = doc.add_paragraph()
            run = p.add_run(f"本周训练主题：{team_diag['training_theme']}")
            run.font.size = Pt(12)
            run.font.name = 'Arial'
            run.bold = True

        if team_diag.get('attention_list'):
            doc.add_heading('关注名单', level=3)
            for item in team_diag['attention_list']:
                add_bullet_paragraph(doc, item, '👀')

        doc.add_paragraph()

    # ============================================================
    # 页脚
    # ============================================================
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('电销录音管理教练 · 数据驱动销售增长 | 录音数据不外传，仅用于团队辅导')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.font.name = 'Arial'

    # ============================================================
    # 保存文件
    # ============================================================
    output_path = data.get('output_path', '')
    if not output_path:
        output_path = os.path.expanduser(f'~/Desktop/录音洞察/{sales_name}/')

    os.makedirs(output_path, exist_ok=True)

    filename = f'{date_str}_{sales_name}_{customer_type}_录音洞察.docx'
    full_path = os.path.join(output_path, filename)

    doc.save(full_path)

    return full_path


def main():
    if len(sys.argv) < 2:
        print('用法: python3 generate_report.py <分析结果.json>', file=sys.stderr)
        print('将分析结果 JSON 生成为格式化的 Word 报告。', file=sys.stderr)
        sys.exit(1)

    input_path = sys.argv[1]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f'错误: 文件不存在 "{input_path}"', file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f'错误: JSON 格式无效 - {e}', file=sys.stderr)
        sys.exit(1)

    output_path = generate_report(data)
    print(json.dumps({
        'success': True,
        'output_path': output_path,
        'sales_name': data.get('sales_name', '未知'),
        'customer_type': data.get('customer_type', '未知'),
    }, ensure_ascii=False))


if __name__ == '__main__':
    main()
